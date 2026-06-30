from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[2]
CONTENT_PATH = Path(__file__).with_name("report_content.md")
OUTPUT_DIR = ROOT / "project" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "JetArm作品设计报告.docx"

SKILL_ROOT = Path(
    r"C:\Users\ASUS\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights  # noqa: E402


BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK_BLUE = "203748"
GRAY = "666666"
LIGHT_FILL = "F4F6F9"
TABLE_BORDER = "C5CED8"
FIGURE_FILL = "F7F9FB"
GOLD = "8A6A16"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, *, color: str = TABLE_BORDER, size: int = 6, style: str = "single") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        node = borders.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), style)
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_run_font(run, *, ascii_font: str = "Calibri", east_asia_font: str = "SimSun", size=None,
                 color: str | None = None, bold: bool | None = None, italic: bool | None = None) -> None:
    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def configure_style_font(style, *, ascii_font: str, east_asia_font: str, size: float,
                         color: str = "000000", bold: bool = False) -> None:
    style.font.name = ascii_font
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.font.bold = bold
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)


def set_outline_level(style, level: int) -> None:
    p_pr = style.element.get_or_add_pPr()
    outline = p_pr.find(qn("w:outlineLvl"))
    if outline is None:
        outline = OxmlElement("w:outlineLvl")
        p_pr.append(outline)
    outline.set(qn("w:val"), str(level))


def configure_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    configure_style_font(normal, ascii_font="Calibri", east_asia_font="SimSun", size=11)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.333
    normal.paragraph_format.first_line_indent = Inches(0.29)

    h1 = doc.styles["Heading 1"]
    configure_style_font(h1, ascii_font="Calibri", east_asia_font="Microsoft YaHei", size=16,
                         color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.page_break_before = True
    set_outline_level(h1, 0)

    h2 = doc.styles["Heading 2"]
    configure_style_font(h2, ascii_font="Calibri", east_asia_font="Microsoft YaHei", size=13,
                         color=BLUE, bold=True)
    h2.paragraph_format.space_before = Pt(12)
    h2.paragraph_format.space_after = Pt(6)
    h2.paragraph_format.keep_with_next = True
    set_outline_level(h2, 1)

    h3 = doc.styles["Heading 3"]
    configure_style_font(h3, ascii_font="Calibri", east_asia_font="Microsoft YaHei", size=12,
                         color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(8)
    h3.paragraph_format.space_after = Pt(4)
    h3.paragraph_format.keep_with_next = True
    set_outline_level(h3, 2)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        configure_style_font(style, ascii_font="Calibri", east_asia_font="SimSun", size=11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.194)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.208

    caption = doc.styles["Caption"]
    configure_style_font(caption, ascii_font="Calibri", east_asia_font="SimSun", size=9.5,
                         color=GRAY, bold=False)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(8)
    caption.paragraph_format.keep_with_next = True

    if "Front Heading" not in doc.styles:
        front = doc.styles.add_style("Front Heading", WD_STYLE_TYPE.PARAGRAPH)
    else:
        front = doc.styles["Front Heading"]
    configure_style_font(front, ascii_font="Calibri", east_asia_font="Microsoft YaHei", size=18,
                         color=INK_BLUE, bold=True)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.space_before = Pt(12)
    front.paragraph_format.space_after = Pt(18)
    front.paragraph_format.keep_with_next = True

    if "Code Block" not in doc.styles:
        code = doc.styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = doc.styles["Code Block"]
    configure_style_font(code, ascii_font="Consolas", east_asia_font="Microsoft YaHei", size=8.5,
                         color="1F2933", bold=False)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.10)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.line_spacing = 1.08

    if "Equation" not in doc.styles:
        eq = doc.styles.add_style("Equation", WD_STYLE_TYPE.PARAGRAPH)
    else:
        eq = doc.styles["Equation"]
    configure_style_font(eq, ascii_font="Cambria Math", east_asia_font="SimSun", size=11,
                         color="000000", bold=False)
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.space_before = Pt(4)
    eq.paragraph_format.space_after = Pt(8)


def configure_page(section) -> None:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True


def add_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("作品设计报告｜云端/本地双模型与RGB视觉闭环机械臂")
    set_run_font(run, east_asia_font="Microsoft YaHei", size=8.5, color=GRAY)

    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "3")
    bottom.set(qn("w:color"), "D9E0E7")
    borders.append(bottom)
    p_pr.append(borders)

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fp.paragraph_format.space_before = Pt(0)
    prefix = fp.add_run("第 ")
    set_run_font(prefix, east_asia_font="SimSun", size=9, color=GRAY)
    add_field(fp, "PAGE", "1")
    suffix = fp.add_run(" 页")
    set_run_font(suffix, east_asia_font="SimSun", size=9, color=GRAY)


def add_cover(doc: Document, lines: list[str]) -> None:
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(12)

    kicker = doc.add_paragraph()
    kicker.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kicker.paragraph_format.space_after = Pt(18)
    run = kicker.add_run("智能机器人作品 · 设计报告")
    set_run_font(run, east_asia_font="Microsoft YaHei", size=11, color=GOLD, bold=True)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    title.paragraph_format.keep_with_next = True
    run = title.add_run(lines[0].lstrip("# ").strip())
    set_run_font(run, east_asia_font="Microsoft YaHei", size=27, color=INK_BLUE, bold=True)

    english = doc.add_paragraph()
    english.alignment = WD_ALIGN_PARAGRAPH.CENTER
    english.paragraph_format.space_after = Pt(24)
    run = english.add_run(lines[2].strip())
    set_run_font(run, east_asia_font="Microsoft YaHei", size=12, color=DARK_BLUE, italic=True)

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    label.paragraph_format.space_after = Pt(54)
    run = label.add_run(lines[4].strip())
    set_run_font(run, east_asia_font="Microsoft YaHei", size=16, color=BLUE, bold=True)

    for raw in lines[6:]:
        if not raw.strip():
            continue
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(9)
        run = p.add_run(raw.strip())
        set_run_font(run, east_asia_font="SimSun", size=12, color="333333")


def add_update_fields_setting(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_after = Pt(4)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "目录将在打开文档时自动更新")


def add_code_block(doc: Document, lines: list[str]) -> None:
    for idx, line in enumerate(lines):
        p = doc.add_paragraph(style="Code Block")
        p.paragraph_format.first_line_indent = Inches(0)
        if idx == 0:
            p.paragraph_format.space_before = Pt(5)
        if idx == len(lines) - 1:
            p.paragraph_format.space_after = Pt(8)
        run = p.add_run(line if line else " ")
        set_run_font(run, ascii_font="Consolas", east_asia_font="Microsoft YaHei", size=8.5,
                     color="1F2933")
        p_pr = p._p.get_or_add_pPr()
        shd = p_pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            p_pr.append(shd)
        shd.set(qn("w:fill"), "F2F4F7")


def clean_md_cell(value: str) -> str:
    return value.strip().replace("\\|", "|")


def parse_table(lines: list[str]) -> list[list[str]]:
    rows = []
    for line in lines:
        cells = [clean_md_cell(v) for v in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", cell) for cell in cells):
            continue
        rows.append(cells)
    return rows


def weighted_widths(rows: list[list[str]]) -> list[int]:
    cols = len(rows[0])
    max_lens = []
    for col in range(cols):
        length = max(len(row[col]) if col < len(row) else 0 for row in rows)
        max_lens.append(max(5, min(length, 36)))
    weights = [max(1.0, length ** 0.62) for length in max_lens]
    if cols >= 7:
        weights = [max(1.0, min(weight, 3.4)) for weight in weights]
    return column_widths_from_weights(weights, 9360)


def add_markdown_table(doc: Document, table_lines: list[str]) -> None:
    rows = parse_table(table_lines)
    if not rows:
        return
    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    widths = weighted_widths(rows)

    for r_idx, values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx in range(cols):
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            value = values[c_idx] if c_idx < len(values) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Inches(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            if cols >= 6 or (r_idx > 0 and len(value) <= 8):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, east_asia_font="Microsoft YaHei" if r_idx == 0 else "SimSun",
                         size=8.5 if cols >= 6 else 9.2, color="1F2933", bold=(r_idx == 0))
            if r_idx == 0:
                set_cell_shading(cell, LIGHT_FILL)
            set_cell_borders(cell)
    apply_table_geometry(
        table,
        widths,
        table_width_dxa=9360,
        indent_dxa=120,
        cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120},
    )
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_figure_placeholder(doc: Document, marker: str) -> None:
    content = marker.strip()[1:-1]
    parts = [part.strip() for part in content.split("|")]
    title = parts[0].replace("FIGURE:", "").strip()
    detail = parts[1].replace("内容:", "").strip() if len(parts) > 1 else ""
    source = parts[2].replace("建议来源:", "").strip() if len(parts) > 2 else ""

    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_shading(cell, FIGURE_FILL)
    set_cell_borders(cell, color="AAB6C2", size=7, style="dashed")

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.space_before = Pt(13)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run(f"【图片占位】{title}")
    set_run_font(run, east_asia_font="Microsoft YaHei", size=10.5, color=DARK_BLUE, bold=True)

    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.first_line_indent = Inches(0)
    p2.paragraph_format.space_after = Pt(4)
    run = p2.add_run(f"图片内容：{detail}")
    set_run_font(run, east_asia_font="SimSun", size=9.3, color="3C4856")

    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p3.paragraph_format.first_line_indent = Inches(0)
    p3.paragraph_format.space_after = Pt(13)
    run = p3.add_run(f"建议来源：{source}。本报告不生成该图片。")
    set_run_font(run, east_asia_font="SimSun", size=9, color=GRAY, italic=True)

    apply_table_geometry(table, [9360], table_width_dxa=9360, indent_dxa=120,
                         cell_margins_dxa={"top": 80, "bottom": 80, "start": 120, "end": 120})
    caption = doc.add_paragraph(title, style="Caption")
    caption.paragraph_format.keep_with_next = False


def is_equation(text: str) -> bool:
    return (
        any(symbol in text for symbol in ("θ", "Δq", "Ji =", "checksum =", "gripper_center =",
                                           "gripper_width =", "gripper_axis =", "X = x", "x = L23", "z = h"))
        and len(text) < 180
    )


def add_body_paragraph(doc: Document, text: str) -> None:
    if is_equation(text):
        p = doc.add_paragraph(style="Equation")
        p.paragraph_format.first_line_indent = Inches(0)
    else:
        p = doc.add_paragraph(style="Normal")
    run = p.add_run(text)
    set_run_font(run, east_asia_font="SimSun", size=11, color="000000")


def render_markdown(doc: Document, lines: list[str]) -> None:
    i = 0
    in_code = False
    code_lines: list[str] = []
    front_headings = {"参赛作品原创性声明", "摘要", "Abstract", "目录"}
    while i < len(lines):
        raw = lines[i].rstrip("\n")
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not stripped:
            i += 1
            continue
        if stripped == "---PAGEBREAK---":
            doc.add_page_break()
            i += 1
            continue
        if stripped == "[[TOC]]":
            add_toc(doc)
            i += 1
            continue
        if stripped.startswith("[FIGURE:"):
            add_figure_placeholder(doc, stripped)
            i += 1
            continue
        if stripped.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1
            add_markdown_table(doc, table_lines)
            continue
        if stripped.startswith("### "):
            p = doc.add_paragraph(stripped[4:].strip(), style="Heading 3")
            p.paragraph_format.first_line_indent = Inches(0)
            i += 1
            continue
        if stripped.startswith("## "):
            p = doc.add_paragraph(stripped[3:].strip(), style="Heading 2")
            p.paragraph_format.first_line_indent = Inches(0)
            i += 1
            continue
        if stripped.startswith("# "):
            title = stripped[2:].strip()
            style = "Front Heading" if title in front_headings else "Heading 1"
            p = doc.add_paragraph(title, style=style)
            p.paragraph_format.first_line_indent = Inches(0)
            i += 1
            continue
        if re.match(r"^-\s+", stripped):
            p = doc.add_paragraph(re.sub(r"^-\s+", "", stripped), style="List Bullet")
            p.paragraph_format.first_line_indent = Inches(-0.194)
            i += 1
            continue
        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
            p.paragraph_format.first_line_indent = Inches(-0.194)
            i += 1
            continue

        add_body_paragraph(doc, stripped)
        i += 1


def build() -> Path:
    content = CONTENT_PATH.read_text(encoding="utf-8")
    lines = content.splitlines()
    first_break = lines.index("---PAGEBREAK---")
    cover_lines = lines[:first_break]
    body_lines = lines[first_break + 1:]

    doc = Document()
    for section in doc.sections:
        configure_page(section)
        configure_header_footer(section)
    configure_styles(doc)

    props = doc.core_properties
    props.title = "基于云端/本地双模型与RGB视觉闭环的桌面机械臂抓放系统"
    props.subject = "参赛作品设计报告"
    props.author = "参赛团队（待填写）"
    props.keywords = "机械臂, RGB视觉, 视觉闭环, 本地模型, 云端模型"
    props.comments = "由项目现有资料生成；未实测指标均已标注。"

    add_cover(doc, cover_lines)
    doc.add_page_break()
    render_markdown(doc, body_lines)
    add_update_fields_setting(doc)

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    return OUTPUT_PATH


if __name__ == "__main__":
    path = build()
    print(path)
