from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Mm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
CONTENT_PATH = Path(__file__).with_name("report_content.md")
OUTPUT_DIR = ROOT / "project" / "deliverables"
OUTPUT_PATH = OUTPUT_DIR / "慧语灵臂_作品设计报告_严格模板版.docx"
CN_TITLE = "慧语灵臂——基于Intel Core Ultra的具身智能机械臂系统"
EN_TITLE = "HUIYU LINGBI: AN EMBODIED INTELLIGENT ROBOTIC ARM SYSTEM BASED ON INTEL CORE ULTRA"

SKILL_ROOT = Path(
    r"C:\Users\ASUS\.codex\plugins\cache\openai-primary-runtime\documents\26.623.12021\skills\documents"
)
sys.path.insert(0, str(SKILL_ROOT / "scripts"))
from table_geometry import apply_table_geometry, column_widths_from_weights, section_content_width_dxa  # noqa: E402


def set_run_font(run, *, west="Times New Roman", east="宋体", size=10.5,
                 bold=None, italic=None, color="000000"):
    run.font.name = west
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), west)
    r_fonts.set(qn("w:hAnsi"), west)
    r_fonts.set(qn("w:eastAsia"), east)


def set_style_font(style, *, west="Times New Roman", east="宋体", size=10.5,
                   bold=False, color="000000"):
    style.font.name = west
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor.from_string(color)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), west)
    r_fonts.set(qn("w:hAnsi"), west)
    r_fonts.set(qn("w:eastAsia"), east)


def set_outline(style, level):
    p_pr = style.element.get_or_add_pPr()
    node = p_pr.find(qn("w:outlineLvl"))
    if node is None:
        node = OxmlElement("w:outlineLvl")
        p_pr.append(node)
    node.set(qn("w:val"), str(level))


def configure_styles(doc):
    normal = doc.styles["Normal"]
    set_style_font(normal, size=10.5)
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.first_line_indent = Pt(21)
    normal.paragraph_format.line_spacing = 1.0
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(0)

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, east="黑体", size=16, bold=True)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.line_spacing = 1.0
    h1.paragraph_format.space_before = Pt(10.5)
    h1.paragraph_format.space_after = Pt(10.5)
    h1.paragraph_format.page_break_before = True
    h1.paragraph_format.keep_with_next = True
    set_outline(h1, 0)

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, east="黑体", size=14, bold=False)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h2.paragraph_format.first_line_indent = Pt(28)
    h2.paragraph_format.line_spacing = 1.0
    h2.paragraph_format.space_before = Pt(5)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True
    set_outline(h2, 1)

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, east="宋体", size=12, bold=False)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    h3.paragraph_format.first_line_indent = Pt(24)
    h3.paragraph_format.line_spacing = 1.0
    h3.paragraph_format.space_before = Pt(3)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.keep_with_next = True
    set_outline(h3, 2)

    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        set_style_font(style, size=10.5)
        style.paragraph_format.left_indent = Pt(21)
        style.paragraph_format.first_line_indent = Pt(-10.5)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.space_after = Pt(0)

    caption = doc.styles["Caption"]
    set_style_font(caption, size=10.5, bold=True)
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.first_line_indent = Pt(0)
    caption.paragraph_format.space_before = Pt(3)
    caption.paragraph_format.space_after = Pt(5)
    caption.paragraph_format.line_spacing = 1.0
    caption.paragraph_format.keep_with_next = False

    code = doc.styles.add_style("模板代码", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(code, west="Consolas", east="宋体", size=9)
    code.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    code.paragraph_format.first_line_indent = Pt(0)
    code.paragraph_format.left_indent = Pt(14)
    code.paragraph_format.line_spacing = 1.0
    code.paragraph_format.space_after = Pt(0)

    eq = doc.styles.add_style("模板公式", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(eq, west="Cambria Math", east="宋体", size=10.5)
    eq.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    eq.paragraph_format.first_line_indent = Pt(0)
    eq.paragraph_format.line_spacing = 1.0
    eq.paragraph_format.space_before = Pt(3)
    eq.paragraph_format.space_after = Pt(3)

    front = doc.styles.add_style("模板前置标题", WD_STYLE_TYPE.PARAGRAPH)
    set_style_font(front, east="黑体", size=16, bold=True)
    front.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    front.paragraph_format.first_line_indent = Pt(0)
    front.paragraph_format.line_spacing = 1.0
    front.paragraph_format.space_before = Pt(10.5)
    front.paragraph_format.space_after = Pt(10.5)


def configure_front_section(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(3.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.different_first_page_header_footer = True


def configure_body_section(section):
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.left_margin = Cm(3.175)
    section.right_margin = Cm(3.175)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.75)
    section.different_first_page_header_footer = False
    section.header.is_linked_to_previous = False
    section.footer.is_linked_to_previous = False


def add_field(paragraph, instruction, placeholder="1"):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = placeholder
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, text, end])


def add_update_fields(doc):
    settings = doc.settings.element
    node = settings.find(qn("w:updateFields"))
    if node is None:
        node = OxmlElement("w:updateFields")
        settings.append(node)
    node.set(qn("w:val"), "true")


def set_header(section, title):
    p = section.header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(title)
    set_run_font(run, east="宋体", size=9)


def set_body_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run("第 ")
    set_run_font(run, east="宋体", size=9)
    add_field(p, "PAGE", "1")
    run = p.add_run(" 页  共 ")
    set_run_font(run, east="宋体", size=9)
    add_field(p, "SECTIONPAGES", "1")
    run = p.add_run(" 页")
    set_run_font(run, east="宋体", size=9)

    sect_pr = section._sectPr
    pg_num = sect_pr.find(qn("w:pgNumType"))
    if pg_num is None:
        pg_num = OxmlElement("w:pgNumType")
        sect_pr.append(pg_num)
    pg_num.set(qn("w:start"), "1")


def add_centered(doc, text, *, east="宋体", west="Times New Roman", size=10.5,
                 bold=False, before=0, after=0):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, west=west, east=east, size=size, bold=bold)
    return p


def add_cover(doc):
    add_centered(doc, "2026年（第十三届）英特尔杯大学生电子设计竞赛嵌入式AI专题赛",
                 east="宋体", size=14)
    add_centered(doc, "2026 Intel Cup Undergraduate Electronic Design Contest",
                 size=14)
    add_centered(doc, "- Embedded System Design Invitational Contest", size=14)
    add_centered(doc, "作品设计报告", east="华文楷体", west="华文楷体", size=32,
                 bold=True, before=39, after=5)
    add_centered(doc, "Final Report", size=16, bold=True, after=10)
    add_centered(doc, "[[TEMPLATE_LOGO]]", size=10.5, after=13)

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(20)
    p.paragraph_format.space_after = Pt(18)
    label = p.add_run("报告题目：")
    set_run_font(label, east="楷体_GB2312", west="楷体_GB2312", size=22)
    value = p.add_run(CN_TITLE)
    set_run_font(value, east="黑体", size=16, bold=True)

    for label_text, value_text in (
        ("学生姓名：", "待填写"),
        ("指导教师：", "待填写"),
        ("参赛学校：", "待填写"),
    ):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(90)
        p.paragraph_format.line_spacing = Pt(25)
        run = p.add_run(label_text)
        set_run_font(run, east="楷体_GB2312", west="楷体_GB2312", size=16)
        run = p.add_run(value_text)
        set_run_font(run, east="楷体_GB2312", west="楷体_GB2312", size=16)


def extract_front(content):
    def between(start, end):
        i = content.index(start) + len(start)
        j = content.index(end, i)
        return content[i:j].strip()

    declaration = between("# 参赛作品原创性声明", "---PAGEBREAK---")
    declaration = [p.strip() for p in declaration.split("\n\n") if p.strip()]
    cn_abstract = between("# 摘要", "关键词：")
    cn_keywords_start = content.index("关键词：")
    cn_keywords_end = content.index("# Abstract", cn_keywords_start)
    cn_keywords = content[cn_keywords_start:cn_keywords_end].strip()
    en_abstract = between("# Abstract", "Keywords:")
    en_keywords_start = content.index("Keywords:")
    en_keywords_end = content.index("---PAGEBREAK---", en_keywords_start)
    en_keywords = content[en_keywords_start:en_keywords_end].strip()
    body = content[content.index("# 第一章 绪论"):]
    return declaration, cn_abstract, cn_keywords, en_abstract, en_keywords, body


def add_declaration(doc, paragraphs):
    doc.add_page_break()
    add_centered(doc, "2026年（第十三届）英特尔杯大学生电子设计竞赛",
                 east="黑体", size=16, bold=True)
    add_centered(doc, "嵌入式AI专题赛", east="黑体", size=16, bold=True, after=13)
    add_centered(doc, "参赛作品原创性声明", east="黑体", size=22, bold=True, after=16)

    body_parts = [p for p in paragraphs if not p.startswith(("参赛成员签字", "指导教师签字", "日期："))]
    for text in body_parts:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Pt(21)
        p.paragraph_format.first_line_indent = Pt(28)
        p.paragraph_format.line_spacing = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, east="宋体", size=14)

    for label in ("参赛队员签名：____________________________",
                  "指导教师签名：____________________________"):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.line_spacing = Pt(24)
        run = p.add_run(label)
        set_run_font(run, east="宋体", size=14)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = Pt(24)
    run = p.add_run("日期：______年____月____日")
    set_run_font(run, east="宋体", size=14)


def add_abstracts(doc, cn_abs, cn_kw, en_abs, en_kw):
    doc.add_page_break()
    add_centered(doc, CN_TITLE, east="黑体", size=16, bold=True, before=10.5, after=10.5)
    add_centered(doc, "摘要", east="黑体", size=14, after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(18)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(cn_abs)
    set_run_font(run, east="宋体", size=10.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    label, value = cn_kw.split("：", 1)
    run = p.add_run(label + "：")
    set_run_font(run, east="黑体", size=12)
    run = p.add_run(value)
    set_run_font(run, east="宋体", size=10.5)

    doc.add_page_break()
    add_centered(doc, EN_TITLE, size=16, bold=True, before=10.5, after=10.5)
    add_centered(doc, "ABSTRACT", size=14, bold=True, after=8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(21)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(en_abs)
    set_run_font(run, size=10.5)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    label, value = en_kw.split(":", 1)
    run = p.add_run(label + ":")
    set_run_font(run, size=12, bold=True)
    run = p.add_run(value)
    set_run_font(run, size=10.5)


def add_toc(doc):
    doc.add_page_break()
    add_centered(doc, "目 录", east="黑体", size=16, before=10.5, after=10.5)
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u', "目录将在Word中自动更新")


def set_cell_margin(cell, top=60, bottom=60, start=80, end=80):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for side, value in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_edges(cell, edges):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        if edge in edges:
            val, size, color = edges[edge]
            node.set(qn("w:val"), val)
            node.set(qn("w:sz"), str(size))
            node.set(qn("w:color"), color)
        else:
            node.set(qn("w:val"), "nil")


def repeat_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    tr_pr.append(node)


def no_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    tr_pr.append(OxmlElement("w:cantSplit"))


def parse_table(lines):
    rows = []
    for line in lines:
        cells = [v.strip() for v in line.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        rows.append(cells)
    return rows


def widths_for(rows, total):
    cols = len(rows[0])
    lens = []
    for c in range(cols):
        length = max(len(r[c]) if c < len(r) else 0 for r in rows)
        lens.append(max(4, min(length, 34)))
    weights = [max(1.0, n ** 0.58) for n in lens]
    return column_widths_from_weights(weights, total)


CAPTIONS = {
    "2.3 性能指标": "系统性能指标",
    "4.2 机械臂关节与功能": "机械臂关节参数",
    "6.2 软件模块划分": "软件模块划分及实现状态",
    "7.2 测试设备": "系统测试设备",
    "7.3 软件单元测试": "软件单元测试结果",
    "7.4 RGB采集性能测试": "RGB图像采集性能记录",
    "7.5 目标检测准确率测试": "目标检测准确率记录",
    "7.6 夹爪定位误差测试": "夹爪定位误差记录",
    "7.7 TCP定位误差测试": "TCP定位误差记录",
    "7.8 视觉闭环收敛时间测试": "视觉闭环收敛时间记录",
    "7.9 J6抓取与释放测试": "J6抓取与释放测试记录",
    "7.10 云端与本地模型对比测试": "云端与本地模型对比",
    "7.11 整机任务与耗时测试": "整机任务耗时记录",
    "7.12 安全测试": "安全功能测试记录",
    "附录A 程序清单": "程序清单",
}


def table_caption(doc, chapter, heading, counter):
    number = f"{chapter}-{counter}" if chapter.isdigit() else f"{chapter}-{counter}"
    title = CAPTIONS.get(heading, heading.split(" ", 1)[-1] if heading else "数据表")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(f"表{number} {title}")
    set_run_font(run, east="宋体", size=10.5, bold=True)


def add_three_line_table(doc, lines, section, chapter, heading, table_counts):
    rows = parse_table(lines)
    if not rows:
        return
    table_counts[chapter] = table_counts.get(chapter, 0) + 1
    table_caption(doc, chapter, heading, table_counts[chapter])

    cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=cols)
    table.autofit = False
    total = section_content_width_dxa(section)
    widths = widths_for(rows, total)
    for ri, values in enumerate(rows):
        row = table.rows[ri]
        no_split(row)
        if ri == 0:
            repeat_header(row)
        for ci in range(cols):
            cell = row.cells[ci]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            value = values[ci] if ci < len(values) else ""
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if cols >= 5 or len(value) <= 10 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(value)
            set_run_font(run, east="宋体", size=10.5, bold=(ri == 0))
            edges = {}
            if ri == 0:
                edges["top"] = ("single", 4, "000000")
                edges["bottom"] = ("single", 4, "000000")
            if ri == len(rows) - 1:
                edges["bottom"] = ("single", 4, "000000")
            set_cell_edges(cell, edges)
            set_cell_margin(cell)
    apply_table_geometry(table, widths, table_width_dxa=total, indent_dxa=108,
                         cell_margins_dxa={"top": 0, "bottom": 0, "start": 108, "end": 108})
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)


def add_figure_box(doc, marker, section):
    content = marker.strip()[1:-1]
    parts = [x.strip() for x in content.split("|")]
    title = parts[0].replace("FIGURE:", "").strip()
    detail = parts[1].replace("内容:", "").strip() if len(parts) > 1 else ""
    source = parts[2].replace("建议来源:", "").strip() if len(parts) > 2 else ""
    total = section_content_width_dxa(section)
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    edges = {e: ("dashed", 8, "7F7F7F") for e in ("top", "left", "bottom", "right")}
    set_cell_edges(cell, edges)
    set_cell_margin(cell, top=100, bottom=100, start=120, end=120)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("【图片位置预留】")
    set_run_font(run, east="黑体", size=10.5, bold=True)
    p = cell.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("图片内容：" + detail)
    set_run_font(run, east="宋体", size=10.5)
    p = cell.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run("建议来源：" + source + "。本报告不生成该图片。")
    set_run_font(run, east="宋体", size=9, italic=True, color="666666")
    apply_table_geometry(table, [total], table_width_dxa=total, indent_dxa=120,
                         cell_margins_dxa={"top": 100, "bottom": 100, "start": 120, "end": 120})
    p = doc.add_paragraph(title, style="Caption")
    p.paragraph_format.keep_with_next = False


def is_equation(text):
    keys = ("θ", "Δq", "Ji =", "checksum =", "gripper_center =", "gripper_width =",
            "gripper_axis =", "X = x", "x = L23", "z = h")
    return len(text) < 180 and any(k in text for k in keys)


def add_code(doc, lines):
    for idx, line in enumerate(lines):
        p = doc.add_paragraph(style="模板代码")
        if idx == 0:
            p.paragraph_format.space_before = Pt(3)
        if idx == len(lines) - 1:
            p.paragraph_format.space_after = Pt(5)
        run = p.add_run(line if line else " ")
        set_run_font(run, west="Consolas", east="宋体", size=9)


def render_body(doc, body, body_section):
    lines = body.splitlines()
    i = 0
    in_code = False
    code_lines = []
    chapter = "1"
    last_heading = ""
    table_counts = {}
    front_titles = {"附录A 程序清单", "附录B 核心源代码与伪代码", "附录C 配置摘要",
                    "附录D 扩展应用系统电路图", "附录E 应用资料与部署说明", "参考文献"}
    while i < len(lines):
        raw = lines[i]
        text = raw.strip()
        if text.startswith("```"):
            if in_code:
                add_code(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_lines.append(raw)
            i += 1
            continue
        if not text or text == "---PAGEBREAK---":
            i += 1
            continue
        if text.startswith("[FIGURE:"):
            add_figure_box(doc, text, body_section)
            i += 1
            continue
        if text.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(lines[i].strip())
                i += 1
            add_three_line_table(doc, block, body_section, chapter, last_heading, table_counts)
            continue
        if text.startswith("### "):
            title = text[4:].strip()
            doc.add_paragraph(title, style="Heading 3")
            last_heading = title
            i += 1
            continue
        if text.startswith("## "):
            title = text[3:].strip()
            doc.add_paragraph(title, style="Heading 2")
            last_heading = title
            i += 1
            continue
        if text.startswith("# "):
            title = text[2:].strip()
            match = re.match(r"第([一二三四五六七八九十]+)章", title)
            if match:
                cn_map = {"一": "1", "二": "2", "三": "3", "四": "4", "五": "5",
                          "六": "6", "七": "7", "八": "8", "九": "9", "十": "10"}
                chapter = cn_map.get(match.group(1), chapter)
            elif title.startswith("附录"):
                chapter = title[2:3]
            elif title == "参考文献":
                chapter = "R"
            doc.add_paragraph(title, style="Heading 1")
            last_heading = title
            i += 1
            continue
        if re.match(r"^-\s+", text):
            doc.add_paragraph(re.sub(r"^-\s+", "", text), style="List Bullet")
            i += 1
            continue
        if re.match(r"^\d+\.\s+", text):
            doc.add_paragraph(re.sub(r"^\d+\.\s+", "", text), style="List Number")
            i += 1
            continue

        style = "模板公式" if is_equation(text) else "Normal"
        p = doc.add_paragraph(style=style)
        if chapter == "R" and re.match(r"^\[\d+\]", text):
            p.paragraph_format.left_indent = Pt(20)
            p.paragraph_format.first_line_indent = Pt(-20)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        set_run_font(run, east="宋体", size=10.5)
        i += 1


def build():
    content = CONTENT_PATH.read_text(encoding="utf-8")
    declaration, cn_abs, cn_kw, en_abs, en_kw, body = extract_front(content)

    doc = Document()
    configure_styles(doc)
    front = doc.sections[0]
    configure_front_section(front)
    set_header(front, CN_TITLE)

    add_cover(doc)
    add_declaration(doc, declaration)
    add_abstracts(doc, cn_abs, cn_kw, en_abs, en_kw)
    add_toc(doc)

    body_section = doc.add_section(WD_SECTION.NEW_PAGE)
    configure_body_section(body_section)
    set_header(body_section, CN_TITLE)
    set_body_footer(body_section)
    render_body(doc, body, body_section)
    add_update_fields(doc)

    props = doc.core_properties
    props.title = CN_TITLE
    props.subject = "2026年英特尔杯大学生电子设计竞赛嵌入式AI专题赛作品设计报告"
    props.author = "参赛团队（待填写）"
    props.keywords = "机械臂, RGB视觉, 视觉闭环, 本地模型, 云端模型"

    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build()
